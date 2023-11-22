import os
import pandas as pd
from docx import Document

# 路径设置
word_folder_path = 'E:\Project\Python\Crawling-around\word批量重命名\word'  # Word文件所在文件夹路径
excel_file_path = 'E:\Project\Python\Crawling-around\word批量重命名\educationdraw.xls'  # Excel文件路径
column_name_excel = 'department_name'  # Excel中包含地区名称的列名
# 初始化一个变量来存储地区信息
city_name = None
county_name = None

# 读取Excel文件
df = pd.read_excel(excel_file_path)
area_to_code = dict(zip(df[column_name_excel], df['department_code']))  # 创建地区名称到编号的映射

# 遍历Word文件夹中的所有文件
for filename in os.listdir(word_folder_path):
    if filename.endswith('.docx'):
        doc_path = os.path.join(word_folder_path, filename)
        doc = Document(doc_path)
        print(f"正在处理{doc_path}...")

        # 初始化变量
        city_name = None
        county_name = None

        # 遍历Word文件夹中的所有文件
        for filename in os.listdir(word_folder_path):
            if filename.endswith('.docx'):
                doc_path = os.path.join(word_folder_path, filename)
                doc = Document(doc_path)
                print(f"正在处理{doc_path}...")

                # 初始化变量
                city_name = None
                county_name = None

                # 获取文件中的第一个表格
                if doc.tables:
                    table = doc.tables[0]
                    for i in range(5, len(table.rows)):
                        city_name = table.cell(i, 1).text.strip()  # 市
                        county_name = table.cell(i, 2).text.strip()  # 县
                        if city_name and county_name:
                            break

                full_area_name = f"{city_name}{county_name}" if city_name and county_name else None
                print("full_area_name:", full_area_name)

                # 如果找到了地区名称并且在Excel文件中有对应的编号
                if full_area_name and full_area_name in area_to_code:
                    base_name = f"{area_to_code[full_area_name]}_{full_area_name}_县级实施工作领导小组名单.docx"
                    new_name = f"{base_name}.docx"
                    new_path = os.path.join(word_folder_path, new_name)
                    # 如果文件名已存在，则添加编号进行区分
                    counter = 1
                    while os.path.exists(new_path):
                        new_name = f"{base_name}_{counter}.docx"
                        new_path = os.path.join(word_folder_path, new_name)
                        counter += 1

                    os.rename(doc_path, new_path)  # 重命名Word文件
                    print(f"文件已重命名为：{new_name}")
                else:
                    print(f"未找到有效的地区信息或Excel中没有对应编号：{filename}")

        print("所有文件重命名完成。")
